#!/usr/bin/env python


def parse_mlst(mlst_tsv_file):
    sample2mlst = {}
    with open(mlst_tsv_file, "r") as f:
        for row in f:
            data = row.rstrip().split("\t")
            # samples/N1097/annotation/N1097.fna
            sample = data[0]
            mlst_scheme = data[1]
            mlst = data[2]
            sample2mlst[sample] = [mlst_scheme, mlst]
    return sample2mlst


def parse_rgi(rgi_file_list, 
              plasmid_contig_list,
              query_cov_cutoff=30,
              aro_filter=[]):
    import pandas
    import re
    import os

    BETALACTAMS = ["monobactam", "carbapenem", "penam", "cephem", "penem", "cephamycin", "cephalosporin", "Beta-Lactam"]
    # carbapenem; cephalosporin; cephamycin; penam
    # create dictionary with sample name as key and list of contigs classified as being on plasmids as values
    sample2plasmid = {}
    for plasmid_file in plasmid_contig_list:
        samplebase = os.path.basename(plasmid_file)
        sample = os.path.splitext(samplebase)[0]
        t = pandas.read_csv(plasmid_file, sep="\t", header=0)
        contig_plasmids = list(t["ID"])
        sample2plasmid[sample] = contig_plasmids

    sample2rgi = {}
    for rgi_file in rgi_file_list:
        print("rgi_file", rgi_file)
        sample = rgi_file.split("/")[1]
        plasmids = sample2plasmid[sample]
        sample2rgi[sample] = {}
        sample2rgi[sample]["transporters"] = []
        sample2rgi[sample]["SNP"] = []
        sample2rgi[sample]["drug_resistance"] = {}
        t = pandas.read_csv(rgi_file, sep="\t", header=0)
        for n, row in t.iterrows():
            if row["Contig"] in plasmids:
                gene = f"{row['Best_hit']} (p)"
            else:
                gene = f"{row['Best_hit']} (c)"

            model_type = row["Model_type"]
            mechanism = row["Mechanism"]
            aro = row["ARO"]
            if aro in aro_filter:
                print("Skipping aro: {aro}")
                continue
            if not isinstance(mechanism, str):
                mechanism = 'n/a'
            coverage = float(row["Percent_coverage"])
            if float(coverage) < query_cov_cutoff:
                print("Skipping low cov entry: %s (%s %%)" % (gene, coverage))
                continue
            identity = float(row["Percent_identity"])
            if "efflux" in mechanism or 'permeability' in mechanism:
                sample2rgi[sample]["transporters"].append([gene, coverage, identity])
                continue

            # protein variant model
            # protein homolog model
            # rRNA gene variant model
            if model_type in ["protein variant model", "rRNA gene variant model"]:
                SNPs_in_Best_Hit_ARO = row["SNPs"]
                sample2rgi[sample]["SNP"].append([gene, SNPs_in_Best_Hit_ARO, coverage, identity])
                continue
            elif model_type == "protein homolog model":
                try:
                    drug_class_list = row["Drug_class"].split("; ")
                except:
                    drug_class_list = ["Unspecified"]
                drug_class_list = list(set([i if i not in BETALACTAMS else "Beta-Lactam" for i in drug_class_list]))

                for drug in drug_class_list:
                    drug = re.sub(" antibiotic", "", drug)
                    if drug in BETALACTAMS:
                        drug = "Beta-Lactam"
                    if drug not in sample2rgi[sample]["drug_resistance"]:
                        sample2rgi[sample]["drug_resistance"][drug] = {}
                    if gene not in sample2rgi[sample]["drug_resistance"][drug]:
                        sample2rgi[sample]["drug_resistance"][drug][gene] = [1, coverage, identity]
                    else:
                        # multiple copies of the same gene, increment count and keep lowest coverage and identity
                        sample2rgi[sample]["drug_resistance"][drug][gene][0] += 1
                        # keep lowest coverage
                        if coverage < sample2rgi[sample]["drug_resistance"][drug][gene][1]:
                            sample2rgi[sample]["drug_resistance"][drug][gene][1] = coverage
                        # keep lowest identity
                        if coverage < sample2rgi[sample]["drug_resistance"][drug][gene][2]:
                            sample2rgi[sample]["drug_resistance"][drug][gene][2] = identity
            else:
                #
                raise IOError("Unknown model type:", model_type)
    return sample2rgi


def parse_mash(mash_file_list):
    sample2species = {}
    for mash_file in mash_file_list:
        sample = mash_file.split("/")[1]
        # 0.999905	998/1000	0	Klebsiella pneumoniae strain 196 map unlocalized plasmid unnamed1 Plasmid_1_Contig_2, whole genome shotgun sequence
        with open(mash_file, "r") as f:
            rows = [i.rstrip().split("\t") for i in f]
            best_hit_score = rows[0][1]
            best_hit_description = ' '.join(rows[0][3].split(" ")[0:2])
            sample2species[sample] = [best_hit_description, best_hit_score]
    return sample2species
    


def generate_report_docx(sample2mlst, 
                         sample2rgi, 
                         sample2species):
    from docx import Document
    from docx.shared import Inches
    document = Document()


    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = f'\t\tLaboratoire de Génomique\n\t\tet de Métagénomique\n\t\tGEN, V.1beta-FR'
    paragraph.style = document.styles["Header"]
    
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = f'Genomics report – analysis number : XX'
    paragraph.style = document.styles["Footer"]
       
    
    document.add_heading('NARA report 1', 0)
    
    document.add_heading('Genome analysis report', level=3)
    
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sample'
    hdr_cells[1].text = 'MLST'

    for sample in sample2mlst:
        row_cells = table.add_row().cells
        row_cells[0].text = sample
        row_cells[1].text = f"{sample2mlst[sample][1]} ({sample2mlst[sample][0]})"

    document.add_page_break()
    document.save('test.docx')


def generate_report_rst(sample2mlst, 
                        sample2rgi, 
                        sample2species,
                        output_name):
    
    import io
    from docutils.core import publish_file, publish_parts
    from docutils.parsers.rst import directives

    print("sample2species within", sample2species)
    ######################
    # MLST table
    ######################
    SAMPLES_LIST = list(sample2species.keys())
    MASH_DATA = []
    for sample in SAMPLES_LIST:
        species = sample2species[sample][0]
        score = sample2species[sample][1]
        MASH_DATA.append(f"{species} ({score})")
    MLST_list = [f"{scheme},{mlst}" for scheme, mlst in sample2mlst.values()]

    table_1_rows = [','.join(i) for i in zip(SAMPLES_LIST, MLST_list, MASH_DATA)]
    table_1 = '\n    '.join(table_1_rows)


    ######################
    # Transporters table
    ######################
    transporters_table_rows = []
    for sample in SAMPLES_LIST:
        if sample in sample2rgi:
            if len(sample2rgi[sample]["transporters"]) != 0:
                transporter_str_list = []
                for tranporter, coverage, identity in sample2rgi[sample]["transporters"]:
                    if identity < 90:
                        transporter_str = f"{tranporter}[1]"
                    else:
                        transporter_str = f"{tranporter}"
                    if coverage < 80:
                        transporter_str += f"[2]"
                    transporter_str_list.append(transporter_str)
                transporters = '; '.join(transporter_str_list)
            else:
                transporters = "No known transporters found"
            transporters_table_rows.append(f"* - {sample}\n      - {transporters}")
        transporters_table = '\n    '.join(transporters_table_rows)


    ######################
    # SNP table
    ######################

    SNP_table_rows = []
    for sample in SAMPLES_LIST:
        if sample in sample2rgi:
            if len(sample2rgi[sample]["SNP"]) != 0:
                SNPs_str_list = []
                for gene, snp, coverage, identity in sample2rgi[sample]["SNP"]:
                    if identity < 90:
                        SNPs_str = f"{gene} ({snp})[1]"
                    else:
                        SNPs_str = f"{gene} ({snp})"
                    if coverage < 80:
                        SNPs_str += f"[2]"
                    SNPs_str_list.append(SNPs_str)
                SNPs = '\n        '.join(SNPs_str_list)
            else:
                SNPs = "No resistance associated SNPs"
            SNP_table_rows.append(f"* - {sample}\n      - {SNPs}")
        SNP_table = '\n    '.join(SNP_table_rows)

    ######################
    # Resistance table
    ######################

    resistance_table_rows = []
    for sample in SAMPLES_LIST:
        if sample in sample2rgi:
            sample_str = ''
            if len(sample2rgi[sample]["drug_resistance"]) != 0:
                drug_list = list(sample2rgi[sample]["drug_resistance"].keys())
                drug_list.sort(key=lambda v: v.upper())
                for drug in drug_list:
                    gene_str_list = []
                    drug_format = f"{drug[0].upper()}{drug[1:len(drug)]}"
                    for gene in sample2rgi[sample]["drug_resistance"][drug]:
                        gene_data = sample2rgi[sample]["drug_resistance"][drug][gene]
                        gene_count, coverage, identity = gene_data
                        # indicate if multiple copies of the same gene
                        if gene_count > 1:
                            gene_label = f'{gene} ({gene_count}x)'
                        else:
                            gene_label = gene
                        if identity < 90:
                            gene_str = f"{gene_label}[1]"
                        else:
                            gene_str = f"{gene_label}"
                        if coverage < 80:
                            gene_str += f"[2]"
                        gene_str_list.append(gene_str)
                    sample_str += f' | **{drug_format}:** ' + ", ".join(gene_str_list) + '\n        '
            else:
                sample_str = "No resistance genes found"
            resistance_table_rows.append(f"* - {sample}\n      - {sample_str}")
        resistance_table = '\n    '.join(resistance_table_rows)
    report_str = f"""

==================
Resistance report
==================

Strain identification 
---------------------

.. csv-table::
    :header: "Sample", "Scheme", "MLST", "Mash best hit"
    :widths: 4, 2, 1, 6

    {table_1}


Antibiotic resistance genes
---------------------------

.. list-table::
   :header-rows: 1
   :widths: 1, 2
   :header-rows: 1

    * - **Sample**
      - **Antibiotic resistance genes**
    {resistance_table}
    
| [1] identity < 90%
| [2] partial gene (<80% of the length of the reference)

Single nucleotide polymorphisms (SNP)
--------------------------------------

.. list-table::
   :header-rows: 1
   :widths: 1, 2
   :header-rows: 1

    * - **Sample**
      - **SNP(s)**
    {SNP_table}

| [1] identity < 90%
| [2] partial gene (<80% of the length of the reference)
 
Antibiotic efflux systems, reduced permeability and regulators
---------------------------------------------------------------

.. list-table::
    :header-rows: 1
    :widths: 1, 2
    
    * - **Sample**
      - **Transporters**
    {transporters_table}

| [1] identity < 90%
| [2] partial gene (<80% of the length of the reference)

    """



    with open(output_name, "w") as fh:
        fh.write(report_str)


rgi_file_list = snakemake.input["rgi_files"]
plasmid_file_list = snakemake.input["plasmid_files"]

aro_filter = snakemake.params["aro_filter"]
# input either list of file or a single file
if not isinstance(rgi_file_list, list):
    rgi_file_list = [rgi_file_list]
mlst_file = snakemake.input["mlst_file"]
mash_file_list = snakemake.input["mash_files"]
# input either list of file or a single file
if not isinstance(mash_file_list, list):
    mash_file_list = [mash_file_list]

output_file = snakemake.output[0]
   
sample2mlst = parse_mlst(mlst_file)
sample2rgi = parse_rgi(rgi_file_list, 
                       plasmid_file_list, 
                       aro_filter=aro_filter)

sample2species = parse_mash(mash_file_list)

sample2mlst = {key:value for key,value in sample2mlst.items() if key in sample2rgi}

generate_report_rst(sample2mlst, 
                    sample2rgi, 
                    sample2species,
                    output_file)
